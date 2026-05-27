from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "HEYMIES_PARTNER_MANUAL.md"
OUTPUT = ROOT / "outputs" / "HeyMies_Partner_Manual.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x64, 0x74, 0x8B)
LIGHT_FILL = "E8EEF5"
BORDER = "CBD5E1"


def set_run_font(run, name="Calibri", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_node = borders.find(qn(f"w:{edge}"))
        if edge_node is None:
            edge_node = OxmlElement(f"w:{edge}")
            borders.append(edge_node)
        edge_node.set(qn("w:val"), "single")
        edge_node.set(qn("w:sz"), "4")
        edge_node.set(qn("w:space"), "0")
        edge_node.set(qn("w:color"), BORDER)

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_field(paragraph, field_instr):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    if "Manual Subtitle" not in styles:
        subtitle = styles.add_style("Manual Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle = styles["Manual Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.paragraph_format.line_spacing = 1.2

    if "Manual Callout" not in styles:
        callout = styles.add_style("Manual Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Manual Callout"]
    callout.font.name = "Calibri"
    callout._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    callout._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    callout.font.size = Pt(11)
    callout.font.color.rgb = INK
    callout.paragraph_format.space_after = Pt(6)
    callout.paragraph_format.line_spacing = 1.25


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("HeyMies Partner Manual")
    set_run_font(run, size=26, bold=True, color=BLUE)

    subtitle = doc.add_paragraph(style="Manual Subtitle")
    subtitle.add_run(
        "A readable guide for product testing, real estate review, buyer journeys, Mia nurture, "
        "lead qualification, and operational checks."
    )

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    set_table_geometry(meta, [2700, 6660])
    rows = [
        ("Audience", "Partner and real estate tester"),
        ("Purpose", "Explain each HeyMies function and provide a practical test path"),
        ("Updated", date.today().strftime("%d %B %Y")),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_FILL)
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=(cell is row.cells[0]), color=INK)

    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    set_table_geometry(callout, [9360])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    paragraph = cell.paragraphs[0]
    paragraph.style = "Manual Callout"
    paragraph.add_run("Core promise: ").bold = True
    paragraph.add_run(
        "By the time an agent or seller receives the lead, they should understand who the buyer is, "
        "what they want, how well the property fits, whether finance and timing look ready, and what "
        "the next action should be."
    )

    doc.add_section(WD_SECTION.NEW_PAGE)


def setup_page(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    for section in doc.sections:
        header = section.header
        header_p = header.paragraphs[0]
        header_p.text = ""
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_p.paragraph_format.space_after = Pt(0)
        run = header_p.add_run("HeyMies Partner Manual")
        set_run_font(run, size=9, color=MUTED)

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.text = ""
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_p.paragraph_format.space_after = Pt(0)
        run = footer_p.add_run("Page ")
        set_run_font(run, size=9, color=MUTED)
        add_field(footer_p, "PAGE")


def add_formatted_text(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=10)
        else:
            paragraph.add_run(part)


def add_body_from_markdown(doc, markdown):
    lines = markdown.splitlines()
    in_intro = True

    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# HeyMies Partner Manual"):
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            marks, text = heading.groups()
            if in_intro:
                in_intro = False
            level = min(len(marks) - 1, 3)
            doc.add_paragraph(text, style=f"Heading {level}")
            continue

        bullet = re.match(r"^-\s+(.+)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, bullet.group(1))
            continue

        number = re.match(r"^\d+\.\s+(.+)$", line)
        if number:
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, number.group(1))
            continue

        p = doc.add_paragraph()
        if line.endswith(":") and len(line) < 80:
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line)
            set_run_font(run, bold=True, color=DARK_BLUE)
        else:
            add_formatted_text(p, line)


def add_table_of_contents_note(doc):
    p = doc.add_paragraph()
    p.style = "Manual Callout"
    run = p.add_run("Navigation note: ")
    set_run_font(run, bold=True, color=DARK_BLUE)
    p.add_run(
        "Use Word's Navigation Pane or update the Table of Contents field after opening the document "
        "if you want generated page numbers. The headings are styled for navigation."
    )


def main():
    doc = Document()
    configure_styles(doc)
    setup_page(doc)
    add_cover(doc)

    doc.add_paragraph("Contents", style="Heading 1")
    add_table_of_contents_note(doc)

    doc.add_paragraph("Manual", style="Heading 1")
    markdown = SOURCE.read_text(encoding="utf-8")
    add_body_from_markdown(doc, markdown)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "HeyMies Partner Manual"
    doc.core_properties.subject = "Partner testing and operations guide"
    doc.core_properties.author = "HeyMies"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
